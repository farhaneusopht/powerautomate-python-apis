import requests
import json
import datetime
import re
import os
from fastapi import FastAPI, Request, Body
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from docx import Document
from pathlib import Path

app = FastAPI()

# ===========================
# Logging
# ===========================
log_file_path = "/app/pad_debug.txt"
os.makedirs("/app", exist_ok=True)

def log(message: str):
    try:
        with open(log_file_path, "a", encoding="utf-8") as f:
            f.write(f"{datetime.datetime.now()} - {message}\n")
    except Exception:
        pass
    print(message)

# ===========================
# Global Exception Handler
# ===========================
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    log(f"[ERROR] Unhandled exception at {request.url}: {repr(exc)}")
    return JSONResponse(
        status_code=500,
        content={"status": "error", "message": str(exc)},
    )

# ===========================
# Helpers
# ===========================
table_row_map = {2: 4, 3: 4, 4: 7, 5: 6, 6: 4, 7: 4, 8: 4, 9: 4,
                 10: 6, 11: 4, 12: 4, 13: 4, 14: 4, 15: 4, 16: 4}
json_file_output = "/app/output_with_bot_replies.json"

def clean_bot_reply(text: str):
    if not text:
        return ""
    text = re.sub(r"(https?://\S+|www\.\S+)", "", text)
    text = re.sub(r"\[.*?\]", "", text)
    return text.strip()

def extract_qa_from_rows(file_path: str, table_row_map: dict, q_col=0, a_col=1):
    try:
        doc = Document(file_path)
    except Exception as e:
        raise Exception(f"Failed to read Word document: {e}")

    if not doc.tables:
        raise Exception("No tables found in Word document")

    results = {}
    for t_index, start_row in table_row_map.items():
        if len(doc.tables) < t_index:
            log(f"Table {t_index} not found in document")
            continue
        table = doc.tables[t_index - 1]
        qa_results = []
        for r_index in range(start_row - 1, len(table.rows)):
            row = table.rows[r_index]
            question = row.cells[q_col].text.strip() if len(row.cells) > q_col else ""
            answer = row.cells[a_col].text.strip() if len(row.cells) > a_col else ""
            if question or answer:
                qa_results.append({"question": question, "answer": answer})
        results[f"Table_{t_index}"] = qa_results
        log(f"Extracted {len(qa_results)} Q/A from Table {t_index} starting row {start_row}")
    return results

# ===========================
# 1. Root
# ===========================
@app.get("/")
def read_root():
    return {"message": "Hello, FastAPI!"}

# ===========================
# 2. Extract QA
# ===========================
@app.post("/extract-qa")
async def extract_qa_api(local_path: str = Body(..., embed=True)):
    """
    JSON body: {"local_path": "/app/uploads/file.docx"}
    """
    if not os.path.exists(local_path):
        return {"status": "error", "message": f"Local file not found: {local_path}"}
    
    qa_map = extract_qa_from_rows(local_path, table_row_map, q_col=0, a_col=1)
    return {"status": "ok", "tables": qa_map, "file_path": local_path}

# ===========================
# 3. Insert Bot Replies
# ===========================
@app.post("/insert-bot-replies")
def insert_bot_replies_api(local_path: str = Body(..., embed=True)):
    """
    JSON body: {"local_path": "/app/uploads/file.docx"}
    """
    if not os.path.exists(local_path):
        return {"status": "error", "message": f"Local file not found: {local_path}"}
    save_path = local_path
    log(f"Using local file at {save_path}")

    # Load JSON data
    if not os.path.exists(json_file_output):
        return {"status": "error", "message": f"Bot replies JSON file not found at {json_file_output}"}
    with open(json_file_output, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Load Word document
    doc = Document(save_path)
    inserted_count = 0
    for table_idx, word_table in enumerate(doc.tables, start=1):
        if table_idx == 1:
            continue
        table_key = f"Table_{table_idx}"
        if table_key not in data["tables"]:
            log(f"Skipping Table {table_idx}, not in JSON")
            continue
        json_rows = data["tables"][table_key]
        if not json_rows:
            continue
        start_row = table_row_map.get(table_idx, 1)
        target_col = 2 if table_idx == 9 else 1
        if len(word_table.rows) < start_row:
            log(f"Skipping Table {table_idx}, not enough rows in Word")
            continue
        for i, word_row in enumerate(word_table.rows[start_row-1:], start=0):
            if i >= len(json_rows):
                break
            bot_reply = clean_bot_reply(json_rows[i].get("bot_reply") or "")
            if len(word_row.cells) <= target_col:
                log(f"Skipping row {i+start_row} in Table {table_idx}, cell index out of range")
                continue
            word_cell = word_row.cells[target_col]
            word_cell.text = ""
            for j, line in enumerate(bot_reply.split("\n")):
                if j == 0:
                    word_cell.text = line.strip()
                else:
                    word_cell.add_paragraph(line.strip())
            inserted_count += 1

    doc.save(save_path)
    log(f"[DONE] Inserted {inserted_count} replies into Word")

    return FileResponse(
        path=save_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=Path(save_path).name
    )

# ===========================
# 4. Start Conversation
# ===========================
class StartConversationRequest(BaseModel):
    secret: str

@app.post("/start-conversation")
def start_conversation(req: StartConversationRequest):
    log(f"[START] /start-conversation")
    try:
        url = "https://directline.botframework.com/v3/directline/conversations"
        headers = {"Authorization": f"Bearer {req.secret}"}
        resp = requests.post(url, headers=headers)
        if resp.status_code == 200:
            return {"status": "ok", "data": resp.json()}
        else:
            return {"status": "error", "code": resp.status_code, "message": resp.text}
    except Exception as e:
        log(f"[ERROR] start_conversation failed: {e}")
        return {"status": "error", "message": str(e)}

# ===========================
# 5. Send Transcript (Word File)
# ===========================
class TranscriptRequest(BaseModel):
    conversation_id: str
    token: str
    local_path: str  # file path in /app/uploads/

def extract_transcript_from_docx(file_path):
    # Word file ko python-docx se open karo
    try:
        doc = Document(file_path)
    except Exception as e:
        raise Exception(f"Failed to read Word document: {e}")

    conversation = []
    msg_id = 1

    # Paragraphs se text extract karo
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            conversation.append({"id": f"msg{msg_id}", "text": text})
            msg_id += 1

    # Tables ke cells ka text bhi extract kar sakte ho (optional)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    conversation.append({"id": f"msg{msg_id}", "text": text})
                    msg_id += 1

    return {"conversation": conversation}
@app.post("/send-transcript")
def send_transcript(req: TranscriptRequest):
    out = {"status": None, "message": None, "sent": []}
    try:
        if not os.path.exists(req.local_path):
            return {"status": "error", "message": f"File not found: {req.local_path}"}

        # Word file se transcript extract
        data = extract_transcript_from_docx(req.local_path)

        headers = {"Authorization": f"Bearer {req.token}"}
        for c in data.get("conversation", []):
            url = f"https://directline.botframework.com/v3/directline/conversations/{req.conversation_id}/activities"
            payload = {"type": "message", "from": {"id": "user"}, "text": c.get("text")}
            resp = requests.post(url, headers=headers, json=payload)
            out["sent"].append({"id": c.get("id"), "status": resp.status_code})

        out["status"] = "ok"
        out["message"] = f"Sent {len(out['sent'])} messages successfully"
        return out

    except Exception as e:
        log(f"[ERROR] send_transcript failed: {e}")
        out["status"] = "error"
        out["message"] = str(e)
        return out
# ===========================
# 6. Get Bot Replies
# ===========================
@app.get("/get-bot-replies")
def get_bot_replies():
    if not os.path.exists(json_file_output):
        return {"status": "error", "message": f"Bot replies JSON file not found at {json_file_output}"}
    try:
        with open(json_file_output, "r", encoding="utf-8") as f:
            data = json.load(f)
        return {"status": "ok", "tables": data.get("tables", {})}
    except Exception as e:
        return {"status": "error", "message": str(e)}
