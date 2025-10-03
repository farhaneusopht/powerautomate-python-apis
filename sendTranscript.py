# import requests
# import json
# import datetime
# import traceback
# from docx import Document  
# import sys
# import time

# def log(msg):
#     with open(r"C:\\Users\\User\\Downloads\\pad_debug.txt", "a", encoding="utf-8") as f:
#         f.write(f"[{datetime.datetime.now()}] {msg}\n")

# # ---- Hardcoded transcript file ----
# # transcript_file = r"C:\Users\User\OneDrive\Documents\Call with Viviana Osborne(External).docx"
# output_file_path = r"C:\Users\User\OneDrive\Desktop\Agent\sendTranscript.json"

# # ---- PAD arguments ----
# if len(sys.argv) < 3:
#     sys.exit(1)
# conversation_id = sys.argv[1]
# token = sys.argv[2]
# transcript_file = sys.argv[3]
# log(f"conversation_id: {conversation_id}")
# log(f"Token: {token}")

# # ---- Read DOCX transcript ----
# doc = Document(transcript_file)
# full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
# words = full_text.split()
# chunk_size = 990  # safer chunk size

# log(f"Total words in transcript: {len(words)}")

# # ---- DirectLine setup ----
# base_url = "https://directline.botframework.com/v3/directline"
# headers = {
#     "Authorization": f"Bearer {token}",
#     "Content-Type": "application/json"
# }

# out = {"status": None, "message": None, "sent": []}
# request_count = 0   # <-- counter for total POST requests

# def send_post_with_retry(url, headers, body, max_retries=3):
#     global request_count
#     for attempt in range(max_retries):
#         try:
#             resp = requests.post(url, headers=headers, json=body, timeout=15)
#             request_count += 1   # <-- count request here
#             return resp
#         except requests.exceptions.RequestException as e:
#             log(f"POST attempt {attempt+1} failed: {e}")
#             time.sleep(2)
#     return None

# def extract_activity_id(resp):
#     """Extract only activity_id"""
#     activity_id = None
#     try:
#         resp_json = resp.json()
#         activity_id = resp_json.get("id")
#     except Exception as e:
#         log(f"Failed to parse activity_id: {e}")
#     return activity_id

# try:
#     # ---- Step 1: Send transcript chunks ----
#     total_chunks = (len(words) + chunk_size - 1) // chunk_size  # ceiling division
#     for i in range(0, len(words), chunk_size):
#         time.sleep(10) # delay between chunks

#         chunk_index = i // chunk_size + 1   # dynamic chunk number
#         chunk_words = words[i:i + chunk_size]
#         chunk_body = " ".join(chunk_words)
#         chunk_text = f"TRANSCRIPT CHUNK {chunk_index}: {chunk_body}"

#         body = {
#             "type": "message",
#             "from": {"id": "user1"},
#             "phase": "transcript_chunk",
#             "chunk_number": chunk_index,
#             "text": chunk_text
#         }

#         resp = send_post_with_retry(
#             f"{base_url}/conversations/{conversation_id}/activities",
#             headers,
#             body
#         )

#         if resp is None:
#             log(f"Chunk {chunk_index} failed after retries, skipping")
#             continue

#         activity_id = extract_activity_id(resp)

#         log(f"Chunk {chunk_index} sent, words={len(chunk_words)}, status={resp.status_code}")
#         out["sent"].append({
#             "chunk_index": chunk_index,
#             "word_count": len(chunk_words),
#             "sent_text": chunk_text,
#             "status": resp.status_code,
#             "response": resp.text,
#             "activity_id": activity_id
#         })

#     # ---- Step 2: Send transcript complete ----
#     time.sleep(20)  # delay before sending COMPLETE
#     complete_text = "TRANSCRIPT COMPLETE"
#     complete_msg = {
#         "type": "message",
#         "from": {"id": "user1"},
#         "phase": "transcript_complete",
#         "total_chunks": total_chunks,   # <-- added info
#         "text": complete_text
#     }
#     resp = send_post_with_retry(
#         f"{base_url}/conversations/{conversation_id}/activities",
#         headers,
#         complete_msg
#     )

#     if resp:
#         activity_id = extract_activity_id(resp)
#         out["sent"].append({
#             "chunk_index": "complete",
#             "sent_text": complete_text,
#             "status": resp.status_code,
#             "response": resp.text,
#             "activity_id": activity_id
#         })
#         log(f"Sent TRANSCRIPT COMPLETE, status={resp.status_code}")
#         out["status"] = "ok"
#         out["message"] = "Transcript phase completed successfully"
#     else:
#         out["status"] = "error"
#         out["message"] = "Failed to send TRANSCRIPT COMPLETE"

#     # ---- Step 3: Final watermark calculation ----
#     final_watermark = request_count * 2
#     out["watermark"] = final_watermark - 1

# except Exception as e:
#     out["status"] = "error"
#     out["message"] = str(e)
#     out["traceback"] = traceback.format_exc()
#     log(f"Error: {e}")

# # ---- Save output ----
# with open(output_file_path, "w", encoding="utf-8") as f:
#     json.dump(out, f, ensure_ascii=False, indent=2)
# print(json.dumps(out, ensure_ascii=False, indent=2))

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse
import requests
import json
import datetime
import traceback
from docx import Document
import time
import os

app = FastAPI()

LOG_FILE = r"C:\Users\User\Downloads\pad_debug.txt"

def log(msg):
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{datetime.datetime.now()}] {msg}\n")

def send_post_with_retry(url, headers, body, max_retries=3):
    for attempt in range(max_retries):
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=15)
            return resp
        except requests.exceptions.RequestException as e:
            log(f"POST attempt {attempt+1} failed: {e}")
            time.sleep(2)
    return None

def extract_activity_id(resp):
    activity_id = None
    try:
        resp_json = resp.json()
        activity_id = resp_json.get("id")
    except Exception as e:
        log(f"Failed to parse activity_id: {e}")
    return activity_id

@app.post("/send/transcript")
async def send_transcript(
    conversation_id: str = Form(...),
    token: str = Form(...),
    transcript_file: UploadFile = File(...)
):
    out = {"status": None, "message": None, "sent": []}
    base_url = "https://directline.botframework.com/v3/directline"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

    try:
        # ---- Read DOCX file ----
        contents = await transcript_file.read()
        tmp_path = f"temp_{transcript_file.filename}"
        with open(tmp_path, "wb") as f:
            f.write(contents)

        doc = Document(tmp_path)
        os.remove(tmp_path)  # clean up

        full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        words = full_text.split()
        chunk_size = 990
        total_chunks = (len(words) + chunk_size - 1) // chunk_size
        log(f"Total words in transcript: {len(words)}")

        request_count = 0

        # ---- Send chunks ----
        for i in range(0, len(words), chunk_size):
            time.sleep(1)  # delay between chunks
            chunk_index = i // chunk_size + 1
            chunk_words = words[i:i + chunk_size]
            chunk_body = " ".join(chunk_words)
            chunk_text = f"TRANSCRIPT CHUNK {chunk_index}: {chunk_body}"

            body = {
                "type": "message",
                "from": {"id": "user1"},
                "phase": "transcript_chunk",
                "chunk_number": chunk_index,
                "text": chunk_text
            }

            resp = send_post_with_retry(
                f"{base_url}/conversations/{conversation_id}/activities",
                headers,
                body
            )
            request_count += 1

            if resp is None:
                log(f"Chunk {chunk_index} failed after retries")
                continue

            activity_id = extract_activity_id(resp)
            out["sent"].append({
                "chunk_index": chunk_index,
                "word_count": len(chunk_words),
                "sent_text": chunk_text,
                "status": resp.status_code,
                "response": resp.text,
                "activity_id": activity_id
            })
            log(f"Chunk {chunk_index} sent, status={resp.status_code}")

        # ---- Send transcript complete ----
        time.sleep(2)
        complete_text = "TRANSCRIPT COMPLETE"
        complete_msg = {
            "type": "message",
            "from": {"id": "user1"},
            "phase": "transcript_complete",
            "total_chunks": total_chunks,
            "text": complete_text
        }

        resp = send_post_with_retry(
            f"{base_url}/conversations/{conversation_id}/activities",
            headers,
            complete_msg
        )
        request_count += 1

        if resp:
            activity_id = extract_activity_id(resp)
            out["sent"].append({
                "chunk_index": "complete",
                "sent_text": complete_text,
                "status": resp.status_code,
                "response": resp.text,
                "activity_id": activity_id
            })
            log(f"Sent TRANSCRIPT COMPLETE, status={resp.status_code}")
            out["status"] = "ok"
            out["message"] = "Transcript phase completed successfully"
        else:
            out["status"] = "error"
            out["message"] = "Failed to send TRANSCRIPT COMPLETE"

        out["watermark"] = request_count * 2 - 1

    except Exception as e:
        out["status"] = "error"
        out["message"] = str(e)
        out["traceback"] = traceback.format_exc()
        log(f"Error: {e}")

    return JSONResponse(content=out)
